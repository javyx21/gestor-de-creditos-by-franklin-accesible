from docx import Document


def export_to_word(title, paragraphs, output_path):
    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(output_path)
